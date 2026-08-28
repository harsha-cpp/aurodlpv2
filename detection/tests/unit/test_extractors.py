"""Extractor coverage across the formats hospitals actually send."""

from __future__ import annotations

import io
import zipfile
from email.message import EmailMessage
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from aurodlpv2_detection.api import detect_email
from aurodlpv2_detection.extractors import extract_bytes
from aurodlpv2_detection.models import Attachment, EmailPayload

PHI_LINE = "Patient Ramesh Kumar Iyer, UHID 0024518, diagnosis E11.9"


def _scan(path: Path, mime: str) -> set[str]:
    result = detect_email(
        EmailPayload(
            attachments=[
                Attachment(
                    id="a1",
                    filename=path.name,
                    mime_type=mime,
                    size_bytes=path.stat().st_size,
                    sha256="x",
                    local_path=str(path),
                )
            ]
        )
    )
    return {entity.type for entity in result.entities}


def test_csv_attachment_is_scanned(tmp_path: Path) -> None:
    """A patient-list CSV is the most common bulk export and was unreadable."""
    path = tmp_path / "patients.csv"
    path.write_text("uhid,name,diagnosis\n0024518,Ramesh Kumar Iyer,E11.9\n", encoding="utf-8")
    assert "MRN" in _scan(path, "text/csv")


def test_plain_text_attachment_is_scanned(tmp_path: Path) -> None:
    path = tmp_path / "note.txt"
    path.write_text(PHI_LINE, encoding="utf-8")
    assert {"MRN", "PERSON"} <= _scan(path, "text/plain")


def test_docx_header_and_textbox_text_is_extracted() -> None:
    """Hospital templates put the patient banner in the header."""
    document = Document()
    document.add_paragraph("Body text with nothing sensitive.")
    document.sections[0].header.paragraphs[0].text = PHI_LINE
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_bytes(
        buffer.getvalue(),
        "report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "a1",
    )
    assert "0024518" in result.text


def test_xlsx_rows_keep_header_context_on_one_line() -> None:
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.append(["UHID", "Name", "Code"])
    sheet.append(["0024518", "Ramesh Kumar Iyer", "E11.9"])
    buffer = io.BytesIO()
    workbook.save(buffer)

    result = extract_bytes(
        buffer.getvalue(),
        "list.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "a1",
    )
    assert "UHID Name Code" in result.text
    assert "0024518" in result.text


def test_eml_attachment_extracts_headers_body_and_nested_files() -> None:
    inner = EmailMessage()
    inner["Subject"] = "Discharge summary"
    inner["From"] = "ward@hospital.in"
    inner.set_content("Please review the attached list.")
    inner.add_attachment(
        b"uhid,name\n0031902,Sunita Bhatt\n",
        maintype="text",
        subtype="csv",
        filename="list.csv",
    )

    result = extract_bytes(bytes(inner), "forwarded.eml", "message/rfc822", "a1")
    assert "Discharge summary" in result.text
    assert "0031902" in result.text


def test_zip_archive_members_are_scanned() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("notes.txt", PHI_LINE)
    result = extract_bytes(buffer.getvalue(), "bundle.zip", "application/zip", "a1")
    assert "0024518" in result.text


def test_archive_depth_is_capped() -> None:
    """A zip inside a zip inside a zip must not recurse without bound."""
    inner = io.BytesIO()
    with zipfile.ZipFile(inner, "w") as archive:
        archive.writestr("notes.txt", PHI_LINE)
    middle = io.BytesIO()
    with zipfile.ZipFile(middle, "w") as archive:
        archive.writestr("inner.zip", inner.getvalue())
    outer = io.BytesIO()
    with zipfile.ZipFile(outer, "w") as archive:
        archive.writestr("middle.zip", middle.getvalue())

    result = extract_bytes(outer.getvalue(), "outer.zip", "application/zip", "a1")
    assert any("depth limit" in error for error in result.errors)


def test_renamed_file_is_still_classified_by_content() -> None:
    """Renaming a spreadsheet to .dat is not a way to skip the scan."""
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.append(["UHID", "0024518"])
    buffer = io.BytesIO()
    workbook.save(buffer)

    result = extract_bytes(buffer.getvalue(), "payload.dat", "application/octet-stream", "a1")
    assert "0024518" in result.text


def test_unsupported_type_reports_an_error_rather_than_silently_passing() -> None:
    result = extract_bytes(b"\x00\x01\x02binary", "thing.bin", "application/octet-stream", "a1")
    assert result.text == ""
    assert result.errors


@pytest.mark.parametrize("extension", [".png", ".jpg"])
def test_images_are_queued_for_ocr(tmp_path: Path, extension: str) -> None:
    path = tmp_path / f"scan{extension}"
    Image.new("RGB", (60, 40), "white").save(path)
    result = extract_bytes(path.read_bytes(), path.name, "application/octet-stream", "a1")
    assert result.ocr_images


def test_ocr_unavailability_is_reported_not_silently_empty() -> None:
    """A scanned page that could not be read must not look like a clean page.

    pytesseract used to live in an optional extra, so a deployment without it
    returned "" for every scanned document and every one of them passed the DLP
    scan. That is a false negative on PHI, not a missing feature.
    """
    from PIL import Image

    from aurodlpv2_detection.config import DetectionConfig
    from aurodlpv2_detection.ocr import extract_text, tesseract_backend

    def _unavailable(*_args: object, **_kwargs: object) -> tuple[str, float]:
        raise tesseract_backend.OcrUnavailableError("pytesseract is not installed")

    original = tesseract_backend.run
    tesseract_backend.run = _unavailable  # type: ignore[assignment]
    try:
        result = extract_text([Image.new("RGB", (40, 20), "white")], DetectionConfig())
    finally:
        tesseract_backend.run = original  # type: ignore[assignment]

    assert result.text == ""
    assert result.errors, "an unreadable page must be reported, not silently skipped"
    assert "ocr unavailable" in result.errors[0]
